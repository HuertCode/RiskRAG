import os
import json
import uuid
from datetime import datetime
from werkzeug.security import generate_password_hash, check_password_hash
from werkzeug.utils import secure_filename
from flask import Flask, render_template, request, redirect, url_for, flash, jsonify, send_file, session, make_response
from flask_sqlalchemy import SQLAlchemy
from flask_login import LoginManager, UserMixin, login_user, login_required, logout_user, current_user
from flask_wtf import FlaskForm
from wtforms import StringField, PasswordField, SubmitField, TextAreaField, FileField, BooleanField
from wtforms.validators import InputRequired, Length, ValidationError
from dotenv import load_dotenv
import PyPDF2
import pdfplumber
import fitz  # PyMuPDF
import docx
import markdown
import bleach
from rag_system import RAGSystem

load_dotenv()

app = Flask(__name__)
app.config['SECRET_KEY'] = os.getenv('SECRET_KEY', 'your-secret-key-here')
app.config['SQLALCHEMY_DATABASE_URI'] = 'sqlite:///risk_kb.db'
app.config['SQLALCHEMY_TRACK_MODIFICATIONS'] = False
app.config['UPLOAD_FOLDER'] = 'uploads'
app.config['MAX_CONTENT_LENGTH'] = 16 * 1024 * 1024  # 16MB max file size

# Create upload directories
os.makedirs(app.config['UPLOAD_FOLDER'], exist_ok=True)
os.makedirs(os.path.join(app.config['UPLOAD_FOLDER'], 'private'), exist_ok=True)
os.makedirs(os.path.join(app.config['UPLOAD_FOLDER'], 'shared'), exist_ok=True)

db = SQLAlchemy(app)
login_manager = LoginManager()
login_manager.init_app(app)
login_manager.login_view = 'login'

# Initialize RAG system
rag_system = RAGSystem()

# Models
class User(UserMixin, db.Model):
    id = db.Column(db.Integer, primary_key=True)
    username = db.Column(db.String(20), unique=True, nullable=False)
    email = db.Column(db.String(120), unique=True, nullable=False)
    password_hash = db.Column(db.String(80), nullable=False)
    is_admin = db.Column(db.Boolean, default=False)
    created_at = db.Column(db.DateTime, default=datetime.utcnow)
    
    def set_password(self, password):
        self.password_hash = generate_password_hash(password)
    
    def check_password(self, password):
        return check_password_hash(self.password_hash, password)

class Document(db.Model):
    id = db.Column(db.Integer, primary_key=True)
    filename = db.Column(db.String(255), nullable=False)
    original_filename = db.Column(db.String(255), nullable=False)
    file_path = db.Column(db.String(500), nullable=False)
    content = db.Column(db.Text)
    is_shared = db.Column(db.Boolean, default=False)
    user_id = db.Column(db.Integer, db.ForeignKey('user.id'), nullable=False)
    created_at = db.Column(db.DateTime, default=datetime.utcnow)
    updated_at = db.Column(db.DateTime, default=datetime.utcnow, onupdate=datetime.utcnow)
    
    user = db.relationship('User', backref=db.backref('documents', lazy=True))

class Note(db.Model):
    id = db.Column(db.Integer, primary_key=True)
    title = db.Column(db.String(255), nullable=False)
    content = db.Column(db.Text, nullable=False)
    is_shared = db.Column(db.Boolean, default=False)
    user_id = db.Column(db.Integer, db.ForeignKey('user.id'), nullable=False)
    created_at = db.Column(db.DateTime, default=datetime.utcnow)
    updated_at = db.Column(db.DateTime, default=datetime.utcnow, onupdate=datetime.utcnow)
    
    user = db.relationship('User', backref=db.backref('notes', lazy=True))

class Chat(db.Model):
    id = db.Column(db.Integer, primary_key=True)
    title = db.Column(db.String(255), nullable=False)
    user_id = db.Column(db.Integer, db.ForeignKey('user.id'), nullable=False)
    created_at = db.Column(db.DateTime, default=datetime.utcnow)
    updated_at = db.Column(db.DateTime, default=datetime.utcnow, onupdate=datetime.utcnow)
    
    user = db.relationship('User', backref=db.backref('chats', lazy=True))
    messages = db.relationship('ChatMessage', backref='chat', lazy=True, cascade='all, delete-orphan')

class ChatMessage(db.Model):
    id = db.Column(db.Integer, primary_key=True)
    chat_id = db.Column(db.Integer, db.ForeignKey('chat.id'), nullable=False)
    content = db.Column(db.Text, nullable=False)
    is_user = db.Column(db.Boolean, default=False)  # True for user message, False for AI response
    citations = db.Column(db.Text)  # JSON string of citations
    created_at = db.Column(db.DateTime, default=datetime.utcnow)

@login_manager.user_loader
def load_user(user_id):
    return User.query.get(int(user_id))

# Forms
class LoginForm(FlaskForm):
    username = StringField('Username', validators=[InputRequired(), Length(min=4, max=20)])
    password = PasswordField('Password', validators=[InputRequired(), Length(min=4, max=80)])
    submit = SubmitField('Login')

class RegisterForm(FlaskForm):
    username = StringField('Username', validators=[InputRequired(), Length(min=4, max=20)])
    email = StringField('Email', validators=[InputRequired(), Length(min=6, max=120)])
    password = PasswordField('Password', validators=[InputRequired(), Length(min=4, max=80)])
    submit = SubmitField('Register')
    
    def validate_username(self, username):
        existing_user = User.query.filter_by(username=username.data).first()
        if existing_user:
            raise ValidationError('Username already exists. Choose a different one.')
    
    def validate_email(self, email):
        existing_user = User.query.filter_by(email=email.data).first()
        if existing_user:
            raise ValidationError('Email already registered. Choose a different one.')

class ResetPasswordForm(FlaskForm):
    email = StringField('Email', validators=[InputRequired(), Length(min=6, max=120)])
    new_password = PasswordField('New Password', validators=[InputRequired(), Length(min=4, max=80)])
    submit = SubmitField('Reset Password')

class NoteForm(FlaskForm):
    title = StringField('Title', validators=[InputRequired(), Length(min=1, max=255)])
    content = TextAreaField('Content', validators=[InputRequired()])
    submit = SubmitField('Save Note')

# Helper functions
def allowed_file(filename):
    ALLOWED_EXTENSIONS = {'txt', 'pdf', 'doc', 'docx', 'md'}
    return '.' in filename and filename.rsplit('.', 1)[1].lower() in ALLOWED_EXTENSIONS

def extract_text_from_file(file_path, filename):
    """Extract text content from uploaded files with improved formatting preservation"""
    text = ""
    file_ext = filename.rsplit('.', 1)[1].lower()
    
    try:
        if file_ext == 'txt' or file_ext == 'md':
            # Try different encodings for text files
            encodings = ['utf-8', 'latin-1', 'cp1252', 'iso-8859-1']
            for encoding in encodings:
                try:
                    with open(file_path, 'r', encoding=encoding) as f:
                        text = f.read()
                        if file_ext == 'md':
                            text = markdown.markdown(text)
                            text = bleach.clean(text, strip=True)
                        break
                except UnicodeDecodeError:
                    continue
            else:
                # If all encodings fail, try with error handling
                with open(file_path, 'r', encoding='utf-8', errors='replace') as f:
                    text = f.read()
                    if file_ext == 'md':
                        text = markdown.markdown(text)
                        text = bleach.clean(text, strip=True)
                        
        elif file_ext == 'pdf':
            # Try multiple PDF extraction methods for better results
            text = extract_pdf_text_improved(file_path)
            
        elif file_ext in ['doc', 'docx']:
            doc = docx.Document(file_path)
            for paragraph in doc.paragraphs:
                if paragraph.text.strip():
                    text += paragraph.text + "\n"
                    
    except Exception as e:
        print(f"Error extracting text from {filename}: {str(e)}")
        text = f"Error reading file: {str(e)}"
    
    # Clean up the extracted text
    text = clean_extracted_text(text)
    return text

def extract_pdf_text_improved(file_path):
    """Extract text from PDF using multiple methods for better results"""
    text = ""
    
    # Method 1: Try pdfplumber (best for preserving formatting)
    try:
        with pdfplumber.open(file_path) as pdf:
            for page in pdf.pages:
                page_text = page.extract_text()
                if page_text:
                    text += page_text + "\n\n"
        if text.strip():
            return text
    except Exception as e:
        print(f"pdfplumber failed: {str(e)}")
    
    # Method 2: Try PyMuPDF (good for complex PDFs)
    try:
        doc = fitz.open(file_path)
        for page in doc:
            page_text = page.get_text()
            if page_text:
                text += page_text + "\n\n"
        doc.close()
        if text.strip():
            return text
    except Exception as e:
        print(f"PyMuPDF failed: {str(e)}")
    
    # Method 3: Fallback to PyPDF2
    try:
        with open(file_path, 'rb') as f:
            pdf_reader = PyPDF2.PdfReader(f)
            for page in pdf_reader.pages:
                page_text = page.extract_text()
                if page_text:
                    text += page_text + "\n"
    except Exception as e:
        print(f"PyPDF2 failed: {str(e)}")
    
    return text

def clean_extracted_text(text):
    """Clean and normalize extracted text"""
    if not text:
        return ""
    
    # Remove excessive whitespace while preserving paragraph breaks
    lines = text.split('\n')
    cleaned_lines = []
    
    for line in lines:
        # Clean each line
        cleaned_line = line.strip()
        
        # Remove common PDF artifacts
        cleaned_line = cleaned_line.replace('\x00', '')  # Null bytes
        cleaned_line = cleaned_line.replace('\x0c', '')  # Form feed
        cleaned_line = cleaned_line.replace('\x0b', '')  # Vertical tab
        
        # Normalize multiple spaces to single space
        cleaned_line = ' '.join(cleaned_line.split())
        
        if cleaned_line:
            cleaned_lines.append(cleaned_line)
        else:
            # Preserve paragraph breaks
            if cleaned_lines and cleaned_lines[-1] != '':
                cleaned_lines.append('')
    
    # Join lines and clean up multiple paragraph breaks
    result = '\n'.join(cleaned_lines)
    
    # Remove excessive blank lines (more than 2 consecutive)
    import re
    result = re.sub(r'\n{3,}', '\n\n', result)
    
    return result.strip()

# Routes
@app.route('/')
def index():
    if current_user.is_authenticated:
        return redirect(url_for('private_documents'))
    return redirect(url_for('login'))

@app.route('/login', methods=['GET', 'POST'])
def login():
    form = LoginForm()
    if form.validate_on_submit():
        user = User.query.filter_by(username=form.username.data).first()
        if user and user.check_password(form.password.data):
            login_user(user)
            return redirect(url_for('private_documents'))
        flash('Invalid username or password')
    return render_template('login.html', form=form)

@app.route('/register', methods=['GET', 'POST'])
def register():
    form = RegisterForm()
    if form.validate_on_submit():
        user = User(username=form.username.data, email=form.email.data)
        user.set_password(form.password.data)
        db.session.add(user)
        db.session.commit()
        flash('Registration successful')
        return redirect(url_for('login'))
    return render_template('register.html', form=form)

@app.route('/reset_password', methods=['GET', 'POST'])
def reset_password():
    form = ResetPasswordForm()
    if form.validate_on_submit():
        user = User.query.filter_by(email=form.email.data).first()
        if user:
            user.set_password(form.new_password.data)
            db.session.commit()
            flash('Password reset successful')
            return redirect(url_for('login'))
        flash('Email not found')
    return render_template('reset_password.html', form=form)

@app.route('/logout')
@login_required
def logout():
    logout_user()
    return redirect(url_for('login'))

@app.route('/private_documents')
@login_required
def private_documents():
    documents = Document.query.filter_by(user_id=current_user.id, is_shared=False).all()
    notes = Note.query.filter_by(user_id=current_user.id, is_shared=False).all()
    return render_template('private_documents.html', documents=documents, notes=notes)

@app.route('/shared_documents')
@login_required
def shared_documents():
    documents = Document.query.filter_by(is_shared=True).all()
    notes = Note.query.filter_by(is_shared=True).all()
    return render_template('shared_documents.html', documents=documents, notes=notes)

@app.route('/ask_me')
@login_required
def ask_me():
    # Get user's chats
    chats = Chat.query.filter_by(user_id=current_user.id).order_by(Chat.updated_at.desc()).all()
    return render_template('ask_me.html', chats=chats)

@app.route('/chat/<int:chat_id>')
@login_required
def get_chat(chat_id):
    """Get chat messages for a specific chat"""
    chat = Chat.query.filter_by(id=chat_id, user_id=current_user.id).first_or_404()
    
    messages = []
    for msg in chat.messages:
        citations = json.loads(msg.citations) if msg.citations else []
        messages.append({
            'id': msg.id,
            'content': msg.content,
            'is_user': msg.is_user,
            'citations': citations,
            'created_at': msg.created_at.isoformat()
        })
    
    return jsonify({
        'chat_id': chat.id,
        'title': chat.title,
        'messages': messages
    })

@app.route('/chat', methods=['POST'])
@login_required
def create_chat():
    """Create a new chat"""
    data = request.get_json()
    title = data.get('title', 'New Chat')
    
    chat = Chat(title=title, user_id=current_user.id)
    db.session.add(chat)
    db.session.commit()
    
    return jsonify({
        'id': chat.id,
        'title': chat.title,
        'created_at': chat.created_at.isoformat()
    })

@app.route('/chat/<int:chat_id>', methods=['DELETE'])
@login_required
def delete_chat(chat_id):
    """Delete a chat"""
    chat = Chat.query.filter_by(id=chat_id, user_id=current_user.id).first_or_404()
    db.session.delete(chat)
    db.session.commit()
    
    return jsonify({'message': 'Chat deleted successfully'})

@app.route('/chat/<int:chat_id>/title', methods=['PUT'])
@login_required
def update_chat_title(chat_id):
    """Update chat title"""
    chat = Chat.query.filter_by(id=chat_id, user_id=current_user.id).first_or_404()
    data = request.get_json()
    title = data.get('title', 'New Chat')
    
    chat.title = title
    chat.updated_at = datetime.utcnow()
    db.session.commit()
    
    return jsonify({'message': 'Chat title updated successfully'})

@app.route('/admin')
@login_required
def admin():
    if not current_user.is_admin:
        flash('Access denied. Admin privileges required.')
        return redirect(url_for('private_documents'))
    
    users = User.query.all()
    
    # Calculate statistics
    total_users = len(users)
    total_admins = sum(1 for user in users if user.is_admin)
    total_documents = sum(len(user.documents) for user in users)
    total_notes = sum(len(user.notes) for user in users)
    
    return render_template('admin.html', 
                         users=users, 
                         total_users=total_users,
                         total_admins=total_admins,
                         total_documents=total_documents,
                         total_notes=total_notes)

@app.route('/upload_document', methods=['POST'])
@login_required
def upload_document():
    if 'file' not in request.files:
        return jsonify({'error': 'No file selected'}), 400
    
    file = request.files['file']
    is_shared = request.form.get('is_shared') == 'true'
    
    if file.filename == '':
        return jsonify({'error': 'No file selected'}), 400
    
    if file and allowed_file(file.filename):
        filename = secure_filename(file.filename)
        unique_filename = f"{uuid.uuid4()}_{filename}"
        
        folder = 'shared' if is_shared else 'private'
        file_path = os.path.join(app.config['UPLOAD_FOLDER'], folder, unique_filename)
        file.save(file_path)
        
        # Extract text content
        content = extract_text_from_file(file_path, filename)
        
        # Save to database
        document = Document(
            filename=unique_filename,
            original_filename=filename,
            file_path=file_path,
            content=content,
            is_shared=is_shared,
            user_id=current_user.id
        )
        db.session.add(document)
        db.session.commit()
        
        # Add to RAG system
        rag_system.add_document(document.id, content, {
            'filename': filename,
            'user_id': current_user.id,
            'is_shared': is_shared,
            'created_at': document.created_at.isoformat()
        })
        
        return jsonify({'message': 'File uploaded successfully'})
    
    return jsonify({'error': 'Invalid file type'}), 400

@app.route('/save_note', methods=['POST'])
@login_required
def save_note():
    data = request.get_json()
    title = data.get('title')
    content = data.get('content')
    is_shared = data.get('is_shared', False)
    note_id = data.get('note_id')
    
    if not title or not content:
        return jsonify({'error': 'Title and content are required'}), 400
    
    if note_id:
        # Update existing note
        note = Note.query.get_or_404(note_id)
        if note.user_id != current_user.id and not current_user.is_admin:
            return jsonify({'error': 'Unauthorized'}), 403
        note.title = title
        note.content = content
        note.is_shared = is_shared
        note.updated_at = datetime.utcnow()
    else:
        # Create new note
        note = Note(
            title=title,
            content=content,
            is_shared=is_shared,
            user_id=current_user.id
        )
        db.session.add(note)
    
    db.session.commit()
    
    # Add/update in RAG system
    rag_system.add_document(f"note_{note.id}", content, {
        'title': title,
        'type': 'note',
        'user_id': current_user.id,
        'is_shared': is_shared,
        'created_at': note.created_at.isoformat()
    })
    
    return jsonify({'message': 'Note saved successfully', 'note_id': note.id})

@app.route('/delete_document/<int:doc_id>', methods=['DELETE'])
@login_required
def delete_document(doc_id):
    document = Document.query.get_or_404(doc_id)
    if document.user_id != current_user.id and not current_user.is_admin:
        return jsonify({'error': 'Unauthorized'}), 403
    
    # Delete file
    if os.path.exists(document.file_path):
        os.remove(document.file_path)
    
    # Remove from RAG system
    rag_system.remove_document(doc_id)
    
    db.session.delete(document)
    db.session.commit()
    
    return jsonify({'message': 'Document deleted successfully'})

@app.route('/delete_note/<int:note_id>', methods=['DELETE'])
@login_required
def delete_note(note_id):
    note = Note.query.get_or_404(note_id)
    if note.user_id != current_user.id and not current_user.is_admin:
        return jsonify({'error': 'Unauthorized'}), 403
    
    # Remove from RAG system
    rag_system.remove_document(f"note_{note_id}")
    
    db.session.delete(note)
    db.session.commit()
    
    return jsonify({'message': 'Note deleted successfully'})

@app.route('/view_document/<int:doc_id>')
@login_required
def view_document(doc_id):
    document = Document.query.get_or_404(doc_id)
    
    # Check permissions
    if not document.is_shared and document.user_id != current_user.id:
        flash('Access denied')
        return redirect(url_for('private_documents'))
    
    return render_template('view_document.html', document=document)

@app.route('/download_document/<int:doc_id>')
@login_required
def download_document(doc_id):
    document = Document.query.get_or_404(doc_id)
    
    # Check permissions
    if not document.is_shared and document.user_id != current_user.id:
        flash('Access denied')
        return redirect(url_for('private_documents'))
    
    # Check if it's a PDF and if inline viewing is requested
    is_pdf = document.original_filename.lower().endswith('.pdf')
    inline = request.args.get('inline', 'false').lower() == 'true'
    
    if is_pdf and inline:
        # Serve PDF for inline viewing
        return send_file(
            document.file_path, 
            mimetype='application/pdf',
            as_attachment=False
        )
    else:
        # Download the file
        return send_file(
            document.file_path, 
            as_attachment=True, 
            download_name=document.original_filename
        )

@app.route('/get_pdf_data/<int:doc_id>')
@login_required
def get_pdf_data(doc_id):
    """Get PDF data as base64 for inline viewing"""
    document = Document.query.get_or_404(doc_id)
    
    # Check permissions
    if not document.is_shared and document.user_id != current_user.id:
        return jsonify({'error': 'Access denied'}), 403
    
    # Check if it's a PDF
    if not document.original_filename.lower().endswith('.pdf'):
        return jsonify({'error': 'Not a PDF file'}), 400
    
    try:
        import base64
        # Read the PDF file and encode as base64
        with open(document.file_path, 'rb') as f:
            pdf_data = f.read()
        
        pdf_base64 = base64.b64encode(pdf_data).decode('utf-8')
        
        return jsonify({
            'pdf_data': f'data:application/pdf;base64,{pdf_base64}',
            'filename': document.original_filename
        })
        
    except Exception as e:
        return jsonify({'error': f'Error reading PDF: {str(e)}'}), 500

@app.route('/ask_question', methods=['POST'])
@login_required
def ask_question():
    data = request.get_json()
    question = data.get('question')
    private_mode = data.get('private_mode', False)
    chat_id = data.get('chat_id')
    
    if not question:
        return jsonify({'error': 'Question is required'}), 400
    
    # Get or create chat
    if chat_id:
        chat = Chat.query.filter_by(id=chat_id, user_id=current_user.id).first()
        if not chat:
            return jsonify({'error': 'Chat not found'}), 404
    else:
        # Create new chat if none specified
        chat = Chat(title="New Chat", user_id=current_user.id)
        db.session.add(chat)
        db.session.commit()
        chat_id = chat.id
    
    # Get relevant documents based on mode
    if private_mode:
        # Only user's private documents
        user_docs = Document.query.filter_by(user_id=current_user.id, is_shared=False).all()
        user_notes = Note.query.filter_by(user_id=current_user.id, is_shared=False).all()
        doc_ids = [doc.id for doc in user_docs] + [f"note_{note.id}" for note in user_notes]
    else:
        # All shared documents
        shared_docs = Document.query.filter_by(is_shared=True).all()
        shared_notes = Note.query.filter_by(is_shared=True).all()
        doc_ids = [doc.id for doc in shared_docs] + [f"note_{note.id}" for note in shared_notes]
    
    # Get answer from RAG system
    answer, citations = rag_system.get_answer(question, doc_ids)
    
    # Save user message
    user_message = ChatMessage(
        chat_id=chat_id,
        content=question,
        is_user=True,
        citations=None
    )
    db.session.add(user_message)
    
    # Save AI response
    ai_message = ChatMessage(
        chat_id=chat_id,
        content=answer,
        is_user=False,
        citations=json.dumps(citations) if citations else None
    )
    db.session.add(ai_message)
    
    # Update chat timestamp
    chat.updated_at = datetime.utcnow()
    
    # Generate or update chat title based on conversation
    new_title = generate_chat_title(chat)
    if new_title and new_title != chat.title:
        chat.title = new_title
    
    db.session.commit()
    
    return jsonify({
        'answer': answer,
        'citations': citations,
        'chat_id': chat_id,
        'chat_title': chat.title
    })

def generate_chat_title(chat):
    """Generate a short, descriptive title for a chat based on its messages"""
    # Get all messages in the chat
    messages = ChatMessage.query.filter_by(chat_id=chat.id).order_by(ChatMessage.created_at.asc()).all()
    
    if not messages:
        return "New Chat"
    
    # Extract user questions (first few messages)
    user_messages = [msg.content for msg in messages if msg.is_user][:3]
    
    if not user_messages:
        return "New Chat"
    
    # Combine user questions for context
    combined_text = " ".join(user_messages)
    
    # Clean and prepare text for summarization
    cleaned_text = clean_text_for_title(combined_text)
    
    # Generate title using different strategies
    title = generate_title_from_text(cleaned_text)
    
    return title

def clean_text_for_title(text):
    """Clean text for title generation"""
    import re
    
    # Remove extra whitespace
    text = re.sub(r'\s+', ' ', text.strip())
    
    # Remove common prefixes
    prefixes_to_remove = [
        'what is', 'what are', 'how to', 'can you', 'could you', 'please',
        'i want to know', 'tell me', 'explain', 'describe', 'show me'
    ]
    
    for prefix in prefixes_to_remove:
        if text.lower().startswith(prefix):
            text = text[len(prefix):].strip()
    
    # Remove punctuation at the end
    text = re.sub(r'[?!.,;:]$', '', text)
    
    return text

def generate_title_from_text(text):
    """Generate a short title from text using multiple strategies"""
    if not text:
        return "New Chat"
    
    # Strategy 1: Use first meaningful sentence (up to 40 chars)
    words = text.split()
    if len(words) <= 6:
        # Short text, use as is (truncated)
        title = text[:40]
        if len(text) > 40:
            title = title.rsplit(' ', 1)[0] + "..."
        return title
    
    # Strategy 2: Extract key topics
    key_topics = extract_key_topics(text)
    if key_topics:
        return key_topics
    
    # Strategy 3: Use first sentence with smart truncation
    first_sentence = text.split('.')[0].strip()
    if len(first_sentence) <= 40:
        return first_sentence
    
    # Strategy 4: Truncate intelligently
    title = text[:40]
    if len(text) > 40:
        # Try to break at word boundary
        last_space = title.rfind(' ')
        if last_space > 25:  # If we can break at a reasonable point
            title = title[:last_space] + "..."
        else:
            title = title + "..."
    
    return title

def extract_key_topics(text):
    """Extract key topics from text for title generation"""
    import re
    
    # Common business/risk topics
    topics = {
        'risk': ['risk', 'risks', 'risk management', 'risk assessment'],
        'compliance': ['compliance', 'regulatory', 'regulation', 'legal'],
        'security': ['security', 'cybersecurity', 'data protection', 'privacy'],
        'finance': ['financial', 'finance', 'budget', 'cost', 'revenue'],
        'operations': ['operational', 'process', 'procedure', 'workflow'],
        'strategy': ['strategy', 'strategic', 'planning', 'goals'],
        'technology': ['technology', 'tech', 'software', 'system', 'platform'],
        'policy': ['policy', 'policies', 'guidelines', 'standards'],
        'audit': ['audit', 'auditing', 'review', 'assessment'],
        'training': ['training', 'education', 'learning', 'development']
    }
    
    text_lower = text.lower()
    
    # Find matching topics
    found_topics = []
    for category, keywords in topics.items():
        for keyword in keywords:
            if keyword in text_lower:
                found_topics.append(category.title())
                break
    
    if found_topics:
        # Combine topics intelligently
        if len(found_topics) == 1:
            return f"{found_topics[0]} Discussion"
        elif len(found_topics) == 2:
            return f"{found_topics[0]} & {found_topics[1]}"
        else:
            return f"{found_topics[0]} & Related Topics"
    
    # If no specific topics found, look for document types
    doc_types = {
        'pdf': ['pdf', 'document', 'report'],
        'policy': ['policy', 'procedure', 'guideline'],
        'report': ['report', 'analysis', 'review'],
        'contract': ['contract', 'agreement', 'terms'],
        'manual': ['manual', 'guide', 'handbook']
    }
    
    for doc_type, keywords in doc_types.items():
        for keyword in keywords:
            if keyword in text_lower:
                return f"{doc_type.title()} Questions"
    
    return None

@app.route('/admin/delete_user/<int:user_id>', methods=['DELETE'])
@login_required
def admin_delete_user(user_id):
    if not current_user.is_admin:
        return jsonify({'error': 'Unauthorized'}), 403
    
    user = User.query.get_or_404(user_id)
    if user.username == 'admin':
        return jsonify({'error': 'Cannot delete admin user'}), 400
    
    # Delete user's documents and notes
    for doc in user.documents:
        if os.path.exists(doc.file_path):
            os.remove(doc.file_path)
        rag_system.remove_document(doc.id)
    
    for note in user.notes:
        rag_system.remove_document(f"note_{note.id}")
    
    db.session.delete(user)
    db.session.commit()
    
    return jsonify({'message': 'User deleted successfully'})

@app.route('/admin/toggle_admin/<int:user_id>', methods=['POST'])
@login_required
def admin_toggle_admin(user_id):
    if not current_user.is_admin:
        return jsonify({'error': 'Unauthorized'}), 403
    
    user = User.query.get_or_404(user_id)
    user.is_admin = not user.is_admin
    db.session.commit()
    
    return jsonify({'message': f'User admin status updated', 'is_admin': user.is_admin})

@app.route('/admin/reprocess_documents', methods=['POST'])
@login_required
def admin_reprocess_documents():
    """Reprocess all documents with improved text extraction"""
    if not current_user.is_admin:
        return jsonify({'error': 'Unauthorized'}), 403
    
    try:
        documents = Document.query.all()
        processed_count = 0
        
        for doc in documents:
            if os.path.exists(doc.file_path):
                # Re-extract text with improved method
                new_content = extract_text_from_file(doc.file_path, doc.original_filename)
                
                # Update document content
                doc.content = new_content
                doc.updated_at = datetime.utcnow()
                
                # Update in RAG system
                rag_system.remove_document(doc.id)
                rag_system.add_document(doc.id, new_content, {
                    'filename': doc.original_filename,
                    'user_id': doc.user_id,
                    'is_shared': doc.is_shared,
                    'created_at': doc.created_at.isoformat()
                })
                
                processed_count += 1
        
        db.session.commit()
        
        return jsonify({
            'message': f'Successfully reprocessed {processed_count} documents with improved text extraction',
            'processed_count': processed_count
        })
        
    except Exception as e:
        return jsonify({'error': f'Error reprocessing documents: {str(e)}'}), 500

@app.route('/get_citation_document/<doc_id>')
@login_required
def get_citation_document(doc_id):
    """Get document information for citation viewing"""
    try:
        # Check if it's a note (doc_id format: "note_123")
        if str(doc_id).startswith('note_'):
            note_id = int(str(doc_id).replace('note_', ''))
            note = Note.query.get(note_id)
            if note:
                # Check permissions
                if not note.is_shared and note.user_id != current_user.id:
                    return jsonify({'error': 'Access denied'}), 403
                
                return jsonify({
                    'type': 'note',
                    'id': note.id,
                    'filename': note.title,
                    'is_pdf': False,
                    'content': note.content,
                    'view_url': None,
                    'download_url': None
                })
        
        # Check if it's a regular document
        try:
            doc_id_int = int(doc_id)
            document = Document.query.get(doc_id_int)
            if document:
                # Check permissions
                if not document.is_shared and document.user_id != current_user.id:
                    return jsonify({'error': 'Access denied'}), 403
                
                return jsonify({
                    'type': 'document',
                    'id': document.id,
                    'filename': document.original_filename,
                    'is_pdf': document.original_filename.lower().endswith('.pdf'),
                    'content': document.content,
                    'view_url': url_for('view_document', doc_id=document.id),
                    'download_url': url_for('download_document', doc_id=document.id)
                })
        except ValueError:
            pass
        
        return jsonify({'error': 'Document not found'}), 404
        
    except Exception as e:
        return jsonify({'error': f'Error retrieving document: {str(e)}'}), 500

if __name__ == '__main__':
    with app.app_context():
        db.create_all()
        
        # Create admin user if not exists
        admin_user = User.query.filter_by(username='admin').first()
        if not admin_user:
            admin_user = User(username='admin', email='admin@risk.com', is_admin=True)
            admin_user.set_password('admin')
            db.session.add(admin_user)
            db.session.commit()
        
        # Reload existing documents into RAG system
        documents = Document.query.all()
        for doc in documents:
            rag_system.add_document(doc.id, doc.content, {
                'filename': doc.original_filename,
                'user_id': doc.user_id,
                'is_shared': doc.is_shared,
                'created_at': doc.created_at.isoformat()
            })
        
        # Reload existing notes into RAG system
        notes = Note.query.all()
        for note in notes:
            rag_system.add_document(f"note_{note.id}", note.content, {
                'title': note.title,
                'type': 'note',
                'user_id': note.user_id,
                'is_shared': note.is_shared,
                'created_at': note.created_at.isoformat()
            })
        
        print(f"Loaded {len(documents)} documents and {len(notes)} notes into RAG system")
    
    app.run(debug=True, host='0.0.0.0', port=5001)